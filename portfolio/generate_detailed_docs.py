from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

doc = Document()

# Title
title = doc.add_heading('Dokumentasi Detail Baris-per-Baris Kode Portofolio', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Penulis: Beni Mulyawan')
doc.add_paragraph('Deskripsi: Dokumen ini membedah kode dari file HTML, CSS, dan JavaScript dari awal hingga akhir beserta fungsinya secara terperinci.')

def add_code_section(title, code, explanation):
    doc.add_heading(title, level=2)
    p_exp = doc.add_paragraph(explanation)
    # Add code block simulation
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Pt(20)
    run = p_code.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ---------------------------------------------------------
doc.add_heading('1. File: index.html (Struktur Utama Halaman)', level=1)

add_code_section('1.1. Deklarasi Dokumen & Head', 
'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>Beni Mulyawan | AI Engineer & Web Developer</title>
  ... (link CSS dan Font) ...
</head>''', 
'Bagian ini adalah standar kerangka HTML5. Tag <meta viewport> memastikan website tampil rapi di HP (responsif). Di sini kita juga memuat font dari Google Fonts (Space Grotesk & Inter) dan menyambungkan file style.css.')

add_code_section('1.2. Background & Efek Kursor', 
'''<body>
  <!-- Latar belakang animasi -->
  <canvas id="bgCanvas"></canvas>
  <div class="scanline"></div>
  
  <!-- Efek kursor custom -->
  <div class="cursor-dot" id="cursorDot"></div>
  <div class="cursor-outline" id="cursorOutline"></div>''', 
'Elemen <canvas> digunakan oleh JavaScript nanti untuk menggambar animasi partikel melayang. <div class="scanline"> memberikan efek visual garis-garis tipis ala monitor jadul/hacker. Bagian cursor digunakan untuk menggantikan kursor mouse bawaan OS dengan kursor bulat bergaya futuristik.')

add_code_section('1.3. Navigasi Bar (Header)', 
'''<header class="header">
  <div class="logo mono">BM_</div>
  <nav class="nav-links">
    <button class="nav-btn is-active" data-target="hero">00/HOME</button>
    <button class="nav-btn" data-target="about">01/ABOUT</button>
    ...
  </nav>
</header>''', 
'Ini adalah menu navigasi di atas. Kita tidak menggunakan tag <a> pindah halaman, melainkan <button> dengan atribut "data-target". Nanti JavaScript akan membaca data-target ini untuk menyembunyikan halaman yang sedang aktif dan memunculkan halaman yang dituju secara instan (Single Page Application).')

add_code_section('1.4. Section Hero (Halaman Utama)', 
'''<section id="hero" class="tab-panel is-active">
  <div class="panel-inner" style="align-items: center; justify-content: center; text-align: center;">
    <p class="hero-subtitle mono animate-item">> SYSTEM READY_</p>
    <h1 class="hero-title animate-item">BENI MULYAWAN</h1>
    ...
</section>''', 
'Ini adalah tampilan pertama saat website dibuka. Diberi class "is-active" agar langsung terlihat. Terdapat efek teks glitch dan animasi teks ngetik (typing) yang dikontrol lewat CSS.')

add_code_section('1.5. Section About & Certifications', 
'''<section id="about" class="tab-panel">
  <div class="panel-inner">
    <!-- View About Utama -->
    <div id="aboutMainView" data-view class="about">
       ... Foto dan deskripsi diri ...
       <button id="btnCertificates">View Certifications</button>
    </div>
    
    <!-- View Sertifikat (Tersembunyi Awalnya) -->
    <div id="certDetailView" data-view style="display: none;">
       <button id="backToAbout">← Back to About</button>
       <div id="certListContainer"></div>
    </div>
  </div>
</section>''', 
'Di bagian About, kita menumpuk dua tampilan sekaligus: aboutMainView (profil) dan certDetailView (sertifikat). View sertifikat awalnya disembunyikan (display: none). Ketika tombol "View Certifications" ditekan, JavaScript akan menyembunyikan profil dan menampilkan sertifikat, membuat pergantian layar yang mulus tanpa menimpa halaman utama.')

add_code_section('1.6. Section Projects & Contact', 
'''<section id="projects" class="tab-panel">
  ... Menampung grid project dan view detail project ...
</section>
<section id="contact" class="tab-panel">
  ... Berisi link sosial media (GitHub, LinkedIn, Email) ...
</section>''', 
'Bagian Projects berfungsi mirip dengan About; memiliki daftar grid proyek dan halaman detail proyek. Bagian Contact berisi tombol-tombol yang mengarah ke link profesional Anda.')


# ---------------------------------------------------------
doc.add_heading('2. File: css/style.css (Desain Tampilan)', level=1)

add_code_section('2.1. CSS Variables (Tema Warna)', 
'''root {
  --bg: #030303;
  --text: #f0f0f0;
  --accent: #ff4d4d;
  --accent-dim: rgba(255, 77, 77, 0.15);
  ...
}''', 
'Di sini kita mendefinisikan variabel warna dasar. Sangat berguna agar warna seragam. --bg adalah warna hitam gelap untuk background, --text adalah putih terang untuk tulisan, dan --accent adalah warna merah cerah untuk tombol dan efek sorotan (hover).')

add_code_section('2.2. Base Styles & Scanline', 
'''body {
  background-color: var(--bg);
  color: var(--text);
  overflow-x: hidden;
}
.scanline {
  background: linear-gradient(to bottom, transparent, rgba(255,255,255,0.02) 50%, transparent);
}''', 
'Mengatur warna seluruh body agar mengambil variabel tema. .scanline adalah efek overlay transparan yang memberikan garis lurus horizontal layaknya monitor CRT lama, menciptakan gaya Cyberpunk/Hacker.')

add_code_section('2.3. Sistem Tab Panel (SPA)', 
'''.tab-panel {
  display: none;
  opacity: 0;
  animation: fadeIn 0.4s ease forwards;
}
.tab-panel.is-active {
  display: block;
}''', 
'Ini adalah rahasia navigasinya. Secara default, semua section disembunyikan (display: none). Tapi jika sebuah section memiliki class "is-active", maka ia akan ditampilkan (display: block) dengan animasi memudar pelan (fadeIn).')

add_code_section('2.4. Desain Grid Proyek & Sertifikat', 
'''.proj-grid, .cert-grid {
  display: grid;
  grid-template-columns: repeat(4, 1fr); /* 4 kolom */
  gap: 1.5rem;
}
@media (max-width: 1024px) { .cert-grid { grid-template-columns: repeat(3, 1fr); } }''', 
'CSS Grid digunakan untuk membagi layar menjadi kotak-kotak. Layar komputer besar mendapat 4 kolom. Aturan @media mengatur agar jika layarnya mengecil (tablet/HP), jumlah kolom otomatis berkurang menjadi 3, 2, atau 1 agar tidak bertumpuk.')


# ---------------------------------------------------------
doc.add_heading('3. File: js/data.js (Pusat Data)', level=1)

add_code_section('3.1. Objek projectsData', 
'''const projectsData = [
  {
    id: "project-1",
    title: "NeuroScan Paper",
    category: "AI/ML",
    image: "images/projects/neuroscan.png",
    description: "Deteksi Dini Tumor Otak Menggunakan Arsitektur CNN MobileNetV2...",
    tech: ["Python", "TensorFlow", "Keras"],
    github: "https://github.com/beni123..."
  },
  ...
];''', 
'Ini adalah database portofolio Anda. Daripada menulis HTML manual berulang-ulang untuk setiap proyek, Anda menyimpannya dalam format Array Object. Nanti main.js akan melooping array ini dan membuat HTML-nya secara otomatis. Kalau mau nambah proyek, cukup tambah kurung kurawal baru di file ini.')


# ---------------------------------------------------------
doc.add_heading('4. File: js/main.js (Logika Interaktif)', level=1)

add_code_section('4.1. Navigasi Antar Tab (SPA Logic)', 
'''const navBtns = document.querySelectorAll('.nav-btn');
const tabPanels = document.querySelectorAll('.tab-panel');

navBtns.forEach(btn => {
  btn.addEventListener('click', () => {
    // Hilangkan semua is-active
    navBtns.forEach(b => b.classList.remove('is-active'));
    tabPanels.forEach(p => p.classList.remove('is-active'));
    
    // Aktifkan yang diklik
    btn.classList.add('is-active');
    const targetId = btn.getAttribute('data-target');
    document.getElementById(targetId).classList.add('is-active');
  });
});''', 
'Ini adalah kode yang merespons ketika menu di header diklik. Ia akan mencari tombol mana yang diklik, membaca "data-target" (misal "projects"), lalu menghapus class "is-active" dari semua halaman, dan memindahkannya hanya ke halaman "projects".')

add_code_section('4.2. Render Proyek dari data.js', 
'''projectsData.forEach(proj => {
  const card = document.createElement('div');
  card.className = 'proj-card';
  card.innerHTML = `<img src="${proj.image}" ...> <h3>${proj.title}</h3>`;
  
  card.addEventListener('click', () => {
    openProjectDetail(proj);
  });
  projGrid.appendChild(card);
});''', 
'Bagian ini membaca data.js tadi. Ia membuat kotak <div> baru untuk setiap proyek, memasukkan foto dan judul ke dalamnya, dan menempelkannya ke layar (projGrid). Ia juga memberikan perintah: "kalau kotak ini diklik, jalankan fungsi openProjectDetail".')

add_code_section('4.3. Logika Grid Sertifikat & Folder', 
'''window.renderCertGrid = function(items, folderName = null) {
  let html = '<div class="cert-grid">';
  items.forEach((item, index) => {
    if (item.type === "folder") {
      html += `<div onclick="openCertFolder(${index})">📁 ${item.name}</div>`;
    } else {
      html += `<a href="${item.file}" target="_blank">📄 ${item.name}</a>`;
    }
  });
  certListContainer.innerHTML = html;
};''', 
'Kode ini bertugas menggambar sertifikat. Ia mengecek: Jika tipe datanya "folder" (seperti Komdigi), ia membuat tombol folder yang kalau diklik akan menjalankan `openCertFolder` untuk menggali ke dalam folder tersebut. Jika bukan folder, ia akan membuat link `<a>` (target="_blank") yang akan membuka file PDF/foto sertifikat di tab browser baru secara langsung.')

add_code_section('4.4. Efek Canvas (Animasi Partikel)', 
'''function initParticles() { ... }
function drawParticles() {
  ctx.clearRect(0, 0, canvas.width, canvas.height);
  particles.forEach(p => {
    p.x += p.vx; p.y += p.vy; // bergerak pelan
    ctx.beginPath();
    ctx.arc(p.x, p.y, p.size, 0, Math.PI * 2);
    ctx.fill();
  });
  requestAnimationFrame(drawParticles);
}''', 
'Ini adalah kode grafis kompleks yang menggambar puluhan titik (partikel) ke layar <canvas>. Setiap titik diberi koordinat x dan y serta kecepatan. Fungsi drawParticles() berjalan terus-menerus 60 frame-per-detik (menggunakan requestAnimationFrame) untuk menggeser posisi titik tersebut, menghasilkan animasi mengambang/melayang di latar belakang web Anda.')

doc.add_paragraph('\n')
p_end = doc.add_paragraph('--- SELESAI ---')
p_end.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.save('C:\\Users\\LENOVO\\Documents\\porto_web_beni\\portfolio\\Penjelasan_Code_Lengkap.docx')
print('Detailed docx generated successfully')
