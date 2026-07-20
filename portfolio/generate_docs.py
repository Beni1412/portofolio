from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Dokumentasi Kode Portofolio Website', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Penulis: Beni Mulyawan')
doc.add_paragraph('Versi: 1.0')
doc.add_paragraph('Deskripsi: Dokumen ini berisi penjelasan lengkap mengenai struktur, kode, dan fitur dari proyek website portofolio.')

# Section 1: Struktur Proyek
doc.add_heading('1. Struktur Proyek', level=1)
doc.add_paragraph('Proyek portofolio ini dibangun menggunakan teknologi native web modern yaitu HTML5, Vanilla CSS, dan Vanilla JavaScript. Struktur folder dirancang sederhana dan modular:')
struktur = [
    ('css/style.css', 'Berisi seluruh desain visual, animasi kursor, transisi halaman, dan grid layout.'),
    ('js/data.js', 'File untuk menyimpan seluruh data konten (metadata) seperti deskripsi project dan tag teknologi.'),
    ('js/main.js', 'File logika inti untuk sistem navigasi, partikel canvas, custom cursor, dan switching antar halaman (project detail, certifications).'),
    ('images/', 'Folder aset gambar dan dokumen sertifikat.'),
    ('index.html', 'Struktur utama dari halaman portofolio, memuat semua section seperti Hero, About, Projects, dan Contact.')
]
for item, desc in struktur:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(item)
    r.bold = True
    p.add_run(f' : {desc}')

# Section 2: Penjelasan Fitur dan Komponen
doc.add_heading('2. Penjelasan Komponen & Logika Inti', level=1)

doc.add_heading('2.1. Navigasi dan Pergantian Halaman (SPA-like Concept)', level=2)
doc.add_paragraph('Portofolio ini menggunakan sistem navigasi Single Page Application (SPA) yang sangat mulus tanpa memuat ulang halaman.')
doc.add_paragraph('Konsepnya memanfaatkan atribut HTML (hidden) atau CSS (display: none / block) untuk menampilkan konten yang aktif. Ketika tombol navigasi ditekan, JavaScript akan mendeteksi atribut data-view dan beralih antar halaman (seperti menukar layar About dengan Certifications).')

doc.add_heading('2.2. Sistem Grid Project dan Sertifikat', level=2)
doc.add_paragraph('Daftar proyek dan sertifikat disajikan dalam tata letak responsif menggunakan CSS Grid. Untuk sertifikat:')
doc.add_paragraph('- Menampilkan 4 kolom berjejer pada layar desktop, yang otomatis menyesuaikan (wrap) ke 3 atau 2 kolom pada perangkat tablet dan handphone.', style='List Bullet')
doc.add_paragraph('- Struktur sertifikat mendukung format Folder. Jika diklik, pengguna akan masuk ke dalam folder dan melihat detail isinya.', style='List Bullet')
doc.add_paragraph('- Apabila pengguna mengklik file (berupa gambar maupun dokumen PDF), sistem akan langsung mengarahkan ke tab baru untuk menampilkan dokumen (menggunakan atribut target="_blank").', style='List Bullet')

doc.add_heading('2.3. Efek Visual (Custom Cursor dan Partikel)', level=2)
doc.add_paragraph('JavaScript digunakan untuk membuat kursor interaktif yang bereaksi terhadap pergerakan mouse dan elemen yang dapat diklik (berubah warna/ukuran). Latar belakang dilengkapi dengan partikel yang dirender pada <canvas> HTML5 untuk memberikan kesan modern dan futuristik.')

# Section 3: Cara Kerja Data Eksternal (data.js)
doc.add_heading('3. Manajemen Data (data.js)', level=1)
doc.add_paragraph('Semua informasi tentang proyek (seperti judul, deskripsi, teknologi yang digunakan, serta URL repositori) disimpan dalam bentuk array objek JSON di dalam file data.js. Hal ini memisahkan antara tampilan (UI) dan data, sehingga untuk menambah proyek baru, cukup menambahkan entri baru di data.js tanpa harus merombak struktur HTML.')

doc.add_paragraph('\n')
p = doc.add_paragraph()
r = p.add_run('--- Akhir dari Dokumentasi ---')
r.italic = True
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.save('C:\\Users\\LENOVO\\Documents\\porto_web_beni\\portfolio\\Dokumentasi_Portfolio_Beni.docx')
print('Docx generated successfully')
